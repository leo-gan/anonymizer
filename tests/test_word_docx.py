"""Word .docx input: load/save, headers, tables, split runs, hyperlinks."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from typer.testing import CliRunner

pytest.importorskip("docx")

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pdf_anonymizer_cli.cli import app
from pdf_anonymizer_core.core import anonymize_docx_file, anonymize_file
from pdf_anonymizer_core.load_and_extract import load_and_extract_text_from_file
from pdf_anonymizer_core.tables import load_review_text
from pdf_anonymizer_core.utils import deanonymize_file, save_results
from pdf_anonymizer_core.word import (
    DOCX_EXTRA_MESSAGE,
    DOCX_SUFFIXES,
    load_docx,
    paragraph_visible_text,
)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    run.append(t_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _build_letter(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Dear Jane Doe,")
    doc.add_paragraph("Email me at jane@acme.com or call 555-123-4567.")
    split = doc.add_paragraph()
    split.add_run("john")
    split.add_run("@")
    split.add_run("example.com")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Email"
    table.cell(1, 0).text = "John Smith"
    table.cell(1, 1).text = "hidden@acme.com"
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Confidential — jane@acme.com"
    footer = section.footer
    footer.paragraphs[0].text = "SSN 123-45-6789"
    link_p = doc.add_paragraph("Contact: ")
    _add_hyperlink(link_p, "jane@acme.com", "mailto:jane@acme.com")
    doc.save(path)
    return path


def _visible_texts(path: Path) -> list[str]:
    loaded = load_docx(str(path))
    return [block.search_text for block in loaded.blocks]


def _anonymize_docx(path: Path, **kwargs):
    defaults = dict(
        characters_to_anonymize=1000,
        prompt_template="unused",
        model_name="unused",
        use_llm=False,
    )
    defaults.update(kwargs)
    return anonymize_docx_file(str(path), **defaults)


class TestDocxLoad:
    def test_reads_body_table_header_footer_and_split_runs(self, tmp_path) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        doc = load_docx(str(src))
        assert doc.kind == "docx"
        texts = [block.search_text for block in doc.blocks]
        assert "Dear Jane Doe," in texts
        assert any("jane@acme.com" in text for text in texts)
        assert "john@example.com" in texts
        assert "hidden@acme.com" in texts
        assert any("Confidential" in text for text in texts)
        assert any("123-45-6789" in text for text in texts)
        assert any(text == "mailto:jane@acme.com" for text in texts)

    def test_flatten_is_not_zip_bytes(self, tmp_path) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        text = load_review_text(str(src))
        assert "# Part:" in text
        assert "jane@acme.com" in text
        assert not text.startswith("PK")


class TestDocxRegexRoundTrip:
    def test_masks_structured_pii_and_preserves_layout(
        self, tmp_path, monkeypatch
    ) -> None:
        monkeypatch.chdir(tmp_path)
        src = _build_letter(tmp_path / "letter.docx")
        review, mapping, entity_texts = _anonymize_docx(src)
        assert "jane@acme.com" not in review
        assert "john@example.com" not in review
        assert "hidden@acme.com" not in review
        assert "123-45-6789" not in review
        assert "EMAIL_" in review
        assert mapping["jane@acme.com"].startswith("EMAIL_")
        assert mapping["123-45-6789"].startswith("SSN")

        save_results(
            review,
            {written: original for original, written in mapping.items()},
            str(src),
            entity_texts=entity_texts,
        )
        out = Path("data/anonymized/letter.anonymized.docx")
        assert out.is_file()
        texts = _visible_texts(out)
        joined = "\n".join(texts)
        assert "jane@acme.com" not in joined
        assert "john@example.com" not in joined
        assert "123-45-6789" not in joined
        assert any(text.startswith("EMAIL_") for text in texts)
        assert "Dear Jane Doe," in texts
        assert "Name" in texts

    def test_split_runs_are_joined_for_detection(self, tmp_path, monkeypatch) -> None:
        monkeypatch.chdir(tmp_path)
        src = tmp_path / "split.docx"
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("ada")
        para.add_run("@")
        para.add_run("example.com")
        doc.save(src)
        review, mapping, entity_texts = _anonymize_docx(src)
        assert "ada@example.com" in mapping
        save_results(
            review,
            {written: original for original, written in mapping.items()},
            str(src),
            entity_texts=entity_texts,
        )
        out = Path("data/anonymized/split.anonymized.docx")
        joined = "\n".join(_visible_texts(out))
        assert "ada@example.com" not in joined
        assert "EMAIL_" in joined

    def test_hyperlink_target_is_rewritten(self, tmp_path, monkeypatch) -> None:
        monkeypatch.chdir(tmp_path)
        src = _build_letter(tmp_path / "letter.docx")
        review, mapping, entity_texts = _anonymize_docx(src)
        save_results(
            review,
            {written: original for original, written in mapping.items()},
            str(src),
            entity_texts=entity_texts,
        )
        out = Path("data/anonymized/letter.anonymized.docx")
        loaded = load_docx(str(out))
        targets = [
            block.search_text for block in loaded.blocks if block.kind == "hyperlink"
        ]
        assert targets
        assert all("jane@acme.com" not in target for target in targets)
        assert any(target.startswith("mailto:EMAIL_") for target in targets)


class TestDocxDeanonymize:
    def test_restores_body_header_and_table(self, tmp_path, monkeypatch) -> None:
        monkeypatch.chdir(tmp_path)
        src = _build_letter(tmp_path / "letter.docx")
        review, mapping, entity_texts = _anonymize_docx(src)
        anon_path, mapping_path = save_results(
            review,
            {written: original for original, written in mapping.items()},
            str(src),
            entity_texts=entity_texts,
        )
        deanonymized_path, _stats = deanonymize_file(anon_path, mapping_path)
        restored = set(_visible_texts(Path(deanonymized_path)))
        source = set(_visible_texts(src))
        assert "jane@acme.com" in "\n".join(restored)
        assert "123-45-6789" in "\n".join(restored)
        assert "hidden@acme.com" in "\n".join(restored)
        assert source == restored


class TestRejectsAndLimits:
    def test_docx_suffix_is_registered(self) -> None:
        assert ".docx" in DOCX_SUFFIXES

    def test_reject_legacy_and_macro_suffixes(self, tmp_path) -> None:
        for name, match in (
            ("letter.doc", "Legacy .doc"),
            ("letter.docm", "Macro-enabled"),
            ("letter.dot", "templates"),
            ("letter.dotm", "Macro-enabled"),
            ("letter.dotx", ".dotx"),
        ):
            path = tmp_path / name
            path.write_bytes(b"not-a-document")
            with pytest.raises(ValueError, match=match):
                anonymize_file(str(path), 1000, "unused", "unused", use_llm=False)
            with pytest.raises(ValueError, match=match):
                load_and_extract_text_from_file(str(path))

    def test_docx_requires_extra(self, tmp_path, monkeypatch) -> None:
        path = tmp_path / "letter.docx"
        path.write_bytes(b"pk")
        monkeypatch.setitem(sys.modules, "docx", None)
        with pytest.raises(ValueError, match=r"pdf-anonymizer-core\[docx\]"):
            load_docx(str(path))
        assert DOCX_EXTRA_MESSAGE in (
            'Word support requires the extra: pip install "pdf-anonymizer-core[docx]"'
        )

    def test_block_limit_raises(self, tmp_path, monkeypatch) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        monkeypatch.setattr("pdf_anonymizer_core.conf.MAX_DOCX_BLOCKS", 2)
        with pytest.raises(ValueError, match="paragraphs"):
            load_docx(str(src))

    def test_byte_limit_raises(self, tmp_path, monkeypatch) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        monkeypatch.setattr("pdf_anonymizer_core.conf.MAX_DOCX_BYTES", 4)
        with pytest.raises(ValueError, match="bytes"):
            load_docx(str(src))

    def test_unreadable_docx_is_value_error(self, tmp_path) -> None:
        path = tmp_path / "bad.docx"
        path.write_bytes(b"not-a-zip")
        with pytest.raises(ValueError, match="Cannot open"):
            load_docx(str(path))

    def test_load_and_extract_refuses_docx(self, tmp_path) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        with pytest.raises(ValueError, match="Word documents"):
            load_and_extract_text_from_file(str(src))


class TestEmptyMappingStillWrites:
    def test_pii_free_docx_is_written(self, tmp_path, monkeypatch) -> None:
        monkeypatch.chdir(tmp_path)
        src = tmp_path / "plain.docx"
        doc = Document()
        doc.add_paragraph("hello world")
        doc.save(src)
        review, mapping, entity_texts = _anonymize_docx(src)
        out, mapping_path = save_results(
            review, mapping, str(src), entity_texts=entity_texts
        )
        assert Path(out).is_file()
        assert Path(mapping_path).is_file()
        assert "hello world" in _visible_texts(Path(out))


class TestAnonymizeFileDispatch:
    def test_docx_dispatch_returns_flatten(self, tmp_path) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        review, mapping = anonymize_file(
            str(src), 1000, "unused", "unused", use_llm=False
        )
        assert review is not None
        assert "# Part:" in review
        assert mapping is not None
        assert "jane@acme.com" in mapping


class TestDocxCli:
    def test_run_writes_docx(self, tmp_path, monkeypatch) -> None:
        monkeypatch.chdir(tmp_path)
        src = _build_letter(tmp_path / "letter.docx")
        runner = CliRunner()
        result = runner.invoke(app, ["run", str(src), "--no-llm"])
        assert result.exit_code == 0, result.output
        out = Path("data/anonymized/letter.anonymized.docx")
        assert out.is_file()
        joined = "\n".join(_visible_texts(out))
        assert "jane@acme.com" not in joined

    def test_verify_rejected_word_suffix_exits_without_traceback(
        self, tmp_path
    ) -> None:
        path = tmp_path / "letter.doc"
        path.write_bytes(b"not-a-document")
        runner = CliRunner()
        result = runner.invoke(app, ["verify", str(path)])
        assert result.exit_code == 1
        assert "Traceback" not in (result.output or "")


class TestDocxLlmFilter:
    def test_llm_hit_must_exist_in_a_block(self, tmp_path, mocker) -> None:
        src = _build_letter(tmp_path / "letter.docx")
        mocker.patch(
            "pdf_anonymizer_core.core.identify_entities_with_llm",
            return_value=[
                {"text": "Jane Doe", "type": "PERSON", "base_form": "Jane Doe"},
                {"text": "Nobody Here", "type": "PERSON", "base_form": "Nobody Here"},
            ],
        )
        mocker.patch(
            "pdf_anonymizer_core.core.extract_entities_via_regex",
            return_value=[],
        )
        _review, mapping, _texts = anonymize_docx_file(
            str(src), 1000, "unused", "unused", use_llm=True
        )
        assert mapping.get("Jane Doe") == "PERSON_1"
        assert "Nobody Here" not in mapping


def test_paragraph_visible_text_skips_deleted_runs(tmp_path) -> None:
    src = tmp_path / "tracked.docx"
    doc = Document()
    para = doc.add_paragraph("keep me")
    deleted = OxmlElement("w:del")
    deleted.set(qn("w:id"), "1")
    deleted.set(qn("w:author"), "Test")
    run = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.text = "secret@acme.com"
    run.append(del_text)
    deleted.append(run)
    para._p.append(deleted)
    doc.save(src)
    loaded = load_docx(str(src))
    texts = [block.search_text for block in loaded.blocks]
    assert "keep me" in texts
    assert not any("secret@acme.com" in text for text in texts)
    assert paragraph_visible_text(para._p) == "keep me"
